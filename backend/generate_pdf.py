#!/usr/bin/env python3
"""
Generate PDF or DOCX version of the Restored Names Bible
Requires: pip install reportlab python-docx
"""

import json
import argparse
from pathlib import Path
import sys
import os

# Add current directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from modernize_language import modernize_bible
except ImportError:
    print("Warning: Could not import modernize_language module")
    modernize_bible = None

def bold_hebrew_names(text):
    """Replace Hebrew names with bold versions for PDF"""
    hebrew_names = [
        'Yahuah', 'Elohiym', 'Yahusha', 'Mashiach', 'Ruach',
        'Qodesh', 'Shaddai', 'Elyon', 'Adonai', 'El',
        'Yahweh', 'YHWH', 'Yah', "Ha'Qodesh", "Ha'Mashiach"
    ]
    
    result = text
    for name in hebrew_names:
        # Case-sensitive replacement with bold tags
        result = result.replace(name, f'<b>{name}</b>')
    
    return result


def generate_pdf(bible_data, output_path, include_toc=True, two_column=False, subtitle=None):
    """Generate a PDF version of the Bible"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
        from reportlab.lib import colors
    except ImportError:
        print("Error: reportlab not installed. Run: pip install reportlab")
        return False

    print(f"Generating PDF: {output_path}")
    
    # Variable to track current book and chapter for headers
    current_book = [None]
    current_chapter = [None]
    
    def add_header_footer(canvas, doc):
        """Add running header with book/chapter and page number"""
        canvas.saveState()
        
        # Header - Running head with book and chapter (only on content pages)
        if current_book[0] and doc.page > 2:  # Skip first few pages (title, TOC)
            canvas.setFont('Helvetica-Oblique', 9)
            canvas.setFillColor(colors.HexColor('#4a5568'))
            
            # Left side: Book name
            canvas.drawString(0.75*inch, letter[1] - 0.5*inch, current_book[0])
            
            # Right side: Chapter number
            if current_chapter[0]:
                chapter_text = f"Chapter {current_chapter[0]}"
                canvas.drawRightString(letter[0] - 0.75*inch, letter[1] - 0.5*inch, chapter_text)
            
            # Draw a subtle line under header
            canvas.setStrokeColor(colors.HexColor('#e2e8f0'))
            canvas.setLineWidth(0.5)
            canvas.line(0.75*inch, letter[1] - 0.55*inch, letter[0] - 0.75*inch, letter[1] - 0.55*inch)
        
        # Footer - Page number (on all pages except title)
        if doc.page > 1:
            canvas.setFont('Helvetica', 9)
            canvas.setFillColor(colors.HexColor('#718096'))
            page_num = f"{doc.page - 1}"  # Just the number
            canvas.drawCentredString(letter[0]/2, 0.5*inch, page_num)
        
        canvas.restoreState()
    
    # Create PDF document
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=1*inch,  # Increased for header
        bottomMargin=0.75*inch,
    )
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Title page style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=36,
        textColor=colors.HexColor('#1a202c'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Times-Bold',
        leading=42
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=16,
        textColor=colors.HexColor('#4a5568'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Times-Italic',
        leading=20
    )
    
    hebrew_names_style = ParagraphStyle(
        'HebrewNames',
        parent=styles['Normal'],
        fontSize=13,
        textColor=colors.HexColor('#2c5282'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Times-Bold',
        leading=18
    )
    
    # Book title style - elegant and prominent
    book_style = ParagraphStyle(
        'BookTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a202c'),
        spaceAfter=30,
        spaceBefore=40,
        alignment=TA_CENTER,
        fontName='Times-Bold',
        leading=28
    )
    
    # Chapter style - subtle and elegant
    chapter_style = ParagraphStyle(
        'ChapterTitle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2d3748'),
        spaceAfter=18,
        spaceBefore=24,
        fontName='Times-Bold',
        leading=20
    )
    
    # Verse style - optimized for readability
    verse_style = ParagraphStyle(
        'VerseText',
        parent=styles['BodyText'],
        fontSize=11,
        leading=18,  # Increased line height
        alignment=TA_JUSTIFY,
        fontName='Times-Roman',
        spaceAfter=4,
        firstLineIndent=0,
        leftIndent=0,
        rightIndent=0
    )
    
    # First verse style with drop cap effect (simulated)
    first_verse_style = ParagraphStyle(
        'FirstVerseText',
        parent=verse_style,
        spaceBefore=6,
        spaceAfter=4
    )
    
    verse_number_style = ParagraphStyle(
        'VerseNumber',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#718096'),
        fontName='Helvetica-Bold'
    )
    
    # Title page - elegant design
    elements.append(Spacer(1, 1.5*inch))
    
    # Decorative line centered
    from reportlab.platypus import HRFlowable
    hr_style = HRFlowable(width="30%", thickness=1, color=colors.HexColor('#2c5282'), 
                          spaceAfter=30, spaceBefore=0, hAlign='CENTER')
    elements.append(hr_style)
    
    elements.append(Paragraph("The Holy Bible", title_style))
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph("King James Version", subtitle_style))
    
    if subtitle:
        elements.append(Paragraph(subtitle, subtitle_style))
    else:
        elements.append(Paragraph("with Restored Hebrew Names", subtitle_style))
        elements.append(Spacer(1, 0.4*inch))
        elements.append(Paragraph("יהוה • אלהים • ישוע • משיח", hebrew_names_style))
        elements.append(Spacer(1, 0.1*inch))
        elements.append(Paragraph("Yahuah • Elohiym • Yahusha • Mashiach", hebrew_names_style))
    
    elements.append(Spacer(1, 0.3*inch))
    # Decorative line
    hr_style2 = HRFlowable(width="30%", thickness=1, color=colors.HexColor('#2c5282'), 
                           spaceAfter=0, spaceBefore=30, hAlign='CENTER')
    elements.append(hr_style2)
    
    elements.append(PageBreak())
    
    # Table of Contents (optional)
    if include_toc:
        elements.append(Paragraph("Table of Contents", book_style))
        elements.append(Spacer(1, 0.3*inch))
        
        toc_data = []
        for book_name in bible_data.keys():
            chapter_count = len(bible_data[book_name])
            toc_data.append([book_name, f"{chapter_count} chapters"])
        
        # Split into Old and New Testament
        ot_books = [
            "Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy",
            "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings",
            "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther",
            "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon",
            "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel",
            "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum",
            "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi"
        ]
        
        elements.append(Paragraph("<b>Old Testament</b>", styles['Heading3']))
        ot_toc = [[book, f"{len(bible_data[book])} chapters"] for book in ot_books if book in bible_data]
        if ot_toc:
            toc_table = Table(ot_toc, colWidths=[3*inch, 1.5*inch])
            toc_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(toc_table)
        
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph("<b>New Testament</b>", styles['Heading3']))
        nt_toc = [[book, f"{len(bible_data[book])} chapters"] for book in bible_data.keys() if book not in ot_books]
        if nt_toc:
            toc_table = Table(nt_toc, colWidths=[3*inch, 1.5*inch])
            toc_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(toc_table)
        
        elements.append(PageBreak())
    
    # Generate content
    for book_name, chapters in bible_data.items():
        print(f"  Processing: {book_name}")
        
        # Update current book for header
        current_book[0] = book_name
        
        # Book title
        elements.append(Paragraph(book_name, book_style))
        elements.append(Spacer(1, 0.2*inch))
        
        for chapter_num in sorted(chapters.keys(), key=int):
            verses = chapters[chapter_num]
            
            # Update current chapter for header
            current_chapter[0] = chapter_num
            
            # Chapter heading with decorative styling
            chapter_heading = f'<font size="16"><b>Chapter {chapter_num}</b></font>'
            elements.append(Paragraph(chapter_heading, chapter_style))
            
            # Add subtle separator after chapter heading
            elements.append(Spacer(1, 0.05*inch))
            
            # Verses
            verse_nums = sorted(verses.keys(), key=int)
            for idx, verse_num in enumerate(verse_nums):
                verse_text = verses[verse_num]
                
                # Bold Hebrew names
                verse_text_formatted = bold_hebrew_names(verse_text)
                
                # Use special style for first verse of chapter (slight emphasis)
                style_to_use = first_verse_style if idx == 0 else verse_style
                
                # Format verse with number (smaller, superscript)
                verse_html = f'<super><font size="7" color="#9ca3af">{verse_num}</font></super> {verse_text_formatted}'
                elements.append(Paragraph(verse_html, style_to_use))
        
        # Page break after each book
        elements.append(PageBreak())
    
    # Build PDF with header/footer
    print("Building PDF document...")
    doc.build(elements, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
    print(f"✓ PDF generated: {output_path}")
    return True


def generate_docx(bible_data, output_path, include_toc=True, subtitle_text=None):
    """Generate a DOCX (Word) version of the Bible"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        return False
    
    def add_verse_with_bold_names(paragraph, verse_text):
        """Add verse text to paragraph with Hebrew names in bold"""
        hebrew_names = [
            'Yahuah', 'Elohiym', 'Yahusha', 'Mashiach', 'Ruach',
            'Qodesh', 'Shaddai', 'Elyon', 'Adonai', 'El',
            'Yahweh', 'YHWH', 'Yah', "Ha'Qodesh", "Ha'Mashiach"
        ]
        
        remaining_text = verse_text
        
        while remaining_text:
            # Find the earliest occurrence of any Hebrew name
            earliest_pos = len(remaining_text)
            earliest_name = None
            
            for name in hebrew_names:
                pos = remaining_text.find(name)
                if pos != -1 and pos < earliest_pos:
                    earliest_pos = pos
                    earliest_name = name
            
            if earliest_name is None:
                # No more names, add remaining text
                run = paragraph.add_run(remaining_text)
                run.font.size = Pt(11)
                break
            else:
                # Add text before the name
                if earliest_pos > 0:
                    run = paragraph.add_run(remaining_text[:earliest_pos])
                    run.font.size = Pt(11)
                
                # Add the name in bold
                run_bold = paragraph.add_run(earliest_name)
                run_bold.font.size = Pt(11)
                run_bold.font.bold = True
                run_bold.font.color.rgb = RGBColor(26, 54, 93)  # Dark blue for emphasis
                
                # Continue with remaining text
                remaining_text = remaining_text[earliest_pos + len(earliest_name):]
    
    print(f"Generating DOCX: {output_path}")
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('The Holy Bible', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('King James Version')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(16)
    subtitle_format.color.rgb = RGBColor(45, 55, 72)
    
    if subtitle_text:
        subtitle2 = doc.add_paragraph(subtitle_text)
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2_format = subtitle2.runs[0].font
        subtitle2_format.size = Pt(14)
        subtitle2_format.color.rgb = RGBColor(45, 55, 72)
    else:
        subtitle2 = doc.add_paragraph('with Restored Hebrew Names')
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2_format = subtitle2.runs[0].font
        subtitle2_format.size = Pt(14)
        subtitle2_format.color.rgb = RGBColor(45, 55, 72)
        
        doc.add_paragraph()
        names = doc.add_paragraph('Yahuah • Elohiym • Yahusha • Mashiach')
        names.alignment = WD_ALIGN_PARAGRAPH.CENTER
        names_format = names.runs[0].font
        names_format.size = Pt(12)
        names_format.italic = True
    
    doc.add_page_break()
    
    # Table of Contents (optional)
    if include_toc:
        toc_heading = doc.add_heading('Table of Contents', 1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        ot_books = [
            "Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy",
            "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings",
            "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther",
            "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon",
            "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel",
            "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum",
            "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi"
        ]
        
        doc.add_heading('Old Testament', 2)
        for book in ot_books:
            if book in bible_data:
                p = doc.add_paragraph(f"{book} — {len(bible_data[book])} chapters")
                p.style = 'List Bullet'
        
        doc.add_heading('New Testament', 2)
        for book in bible_data.keys():
            if book not in ot_books:
                p = doc.add_paragraph(f"{book} — {len(bible_data[book])} chapters")
                p.style = 'List Bullet'
        
        doc.add_page_break()
    
    # Generate content
    for book_idx, (book_name, chapters) in enumerate(bible_data.items()):
        print(f"  Processing: {book_name}")
        
        # Add section for each book (allows different headers)
        if book_idx > 0:
            doc.add_section()
        
        # Set header for this section
        section = doc.sections[-1]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = book_name
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.runs[0] if header_para.runs else header_para.add_run(book_name)
        header_run.font.size = Pt(10)
        header_run.font.italic = True
        header_run.font.color.rgb = RGBColor(113, 128, 150)
        
        # Set footer with page number
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run()
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(113, 128, 150)
        
        # Add page number field
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        footer_run._r.append(fldChar1)
        footer_run._r.append(instrText)
        footer_run._r.append(fldChar2)
        
        # Book title
        book_heading = doc.add_heading(book_name, 1)
        book_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for chapter_num in sorted(chapters.keys(), key=int):
            verses = chapters[chapter_num]
            
            # Chapter heading
            doc.add_heading(f'Chapter {chapter_num}', 2)
            
            # Verses
            for verse_num in sorted(verses.keys(), key=int):
                verse_text = verses[verse_num]
                
                # Create paragraph with verse number and text
                p = doc.add_paragraph()
                
                # Verse number (superscript)
                run_num = p.add_run(verse_num)
                run_num.font.superscript = True
                run_num.font.bold = True
                run_num.font.size = Pt(9)
                run_num.font.color.rgb = RGBColor(113, 128, 150)
                
                # Space
                p.add_run(' ')
                
                # Verse text with bold Hebrew names
                add_verse_with_bold_names(p, verse_text)
        
        # Don't add page break after last book
        if book_idx < len(bible_data) - 1:
            doc.add_page_break()
    
    # Save document
    print("Saving DOCX document...")
    doc.save(str(output_path))
    print(f"✓ DOCX generated: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Generate PDF or DOCX version of the Restored Names Bible'
    )
    parser.add_argument(
        '--input',
        default='build/restored_kjv.json',
        help='Input JSON file (default: build/restored_kjv.json)'
    )
    parser.add_argument(
        '--output',
        help='Output file path (default: build/restored_kjv.pdf or .docx)'
    )
    parser.add_argument(
        '--format',
        choices=['pdf', 'docx', 'both'],
        default='pdf',
        help='Output format (default: pdf)'
    )
    parser.add_argument(
        '--no-toc',
        action='store_true',
        help='Skip table of contents'
    )
    parser.add_argument(
        '--modernize',
        action='store_true',
        help='Modernize archaic language (thee, thou, etc.)'
    )
    parser.add_argument(
        '--all-versions',
        action='store_true',
        help='Generate both original and modernized versions'
    )
    
    args = parser.parse_args()
    
    # Load Bible data
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        return 1
    
    print(f"Loading Bible data from: {input_path}")
    with open(input_path, 'r', encoding='utf-8') as f:
        bible_data = json.load(f)
    
    print(f"Loaded {len(bible_data)} books")
    
    include_toc = not args.no_toc
    success = True
    
    # Prepare versions to generate
    versions = []
    
    if args.all_versions:
        versions = [
            ('original', bible_data, None),
            ('modernized', None, 'with Restored Names & Modernized Language')
        ]
    elif args.modernize:
        versions = [('modernized', None, 'with Restored Names & Modernized Language')]
    else:
        versions = [('original', bible_data, None)]
    
    # Process each version
    for version_name, version_data, subtitle in versions:
        # Apply modernization if needed
        if version_data is None:
            if modernize_bible is None:
                print("Error: modernize_language module not available")
                return 1
            
            print(f"\nApplying language modernization...")
            replacements = {
                " thee ": " you ", " thou ": " you ", " thy ": " your ", " thine ": " your ",
                " ye ": " you ", " art ": " are ", " hast ": " have ", " hadst ": " had ",
                " doest ": " do ", " didst ": " did ", " wilt ": " will ", " shalt ": " shall ",
                " shouldst ": " should ", " wouldst ": " would ", " mayest ": " may ",
                " mightest ": " might ", " canst ": " can ", " couldst ": " could ",
                " knowest ": " know ", " sayest ": " say ", " saith ": " says ",
                " doth ": " does ", " hath ": " has ", " spake ": " spoke ",
                " shew ": " show ", " shewed ": " showed ", " betwixt ": " between ",
                " unto ": " to ",
            }
            version_data = modernize_bible(bible_data, replacements)
        
        # Generate PDF
        if args.format in ['pdf', 'both']:
            if args.output:
                output_path = Path(args.output)
            else:
                suffix = '_modernized' if version_name == 'modernized' else ''
                output_path = Path(f'build/restored_kjv_bible{suffix}.pdf')
            
            success = generate_pdf(version_data, output_path, include_toc=include_toc, subtitle=subtitle) and success
        
        # Generate DOCX
        if args.format in ['docx', 'both']:
            if args.output and args.format != 'both':
                output_path = Path(args.output)
            else:
                suffix = '_modernized' if version_name == 'modernized' else ''
                output_path = Path(f'build/restored_kjv_bible{suffix}.docx')
            
            success = generate_docx(version_data, output_path, include_toc=include_toc, subtitle_text=subtitle) and success
    
    if args.all_versions:
        print("\n✓ Generated both original and modernized versions")
    
    return 0 if success else 1


if __name__ == '__main__':
    exit(main())

